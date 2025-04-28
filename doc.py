from docx import Document

# Criar o documento
doc = Document()
doc.add_heading('Relatório de Teste: SigleClient x MultiClient', 0)

# Adicionar configuração
doc.add_heading('Configuração de Teste:', level=1)
doc.add_paragraph('- 1 único cliente x 10 clientes simultâneos')
doc.add_paragraph('- 100 requisições')
doc.add_paragraph('- Máquina: i3 com 10G de RAM')

doc.add_heading('SigleClient - MultiClient', level=2)

table = doc.add_table(rows=1, cols=5)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Servidor'
hdr_cells[1].text = 'Status'
hdr_cells[2].text = 'Tempo (s)'
hdr_cells[3].text = 'CPU (%)'
hdr_cells[4].text = 'Memória (MB)'

# Preencher dados
data = [
    ('OneShotServer', 'Sucesso - Falhou',  "1.8s - 1.8s", f'1{"%"} - 2%', '9.6MB - 11MB'),
    ('ThreadedServer', 'Sucesso - Sucesso', "1.6s - 3.2s",  f'1.9{"%"} -13%', '10MB - 11MB'),
    ('ThreadPoolServer', 'Sucesso - Sucesso', "21.5s - 16.3s",  f'1{"%"} - 6%', '11MB - 11MB'),
]

for servidor, resultado, tempo, cpu, memoria in data:
    row_cells = table.add_row().cells
    row_cells[0].text = servidor
    row_cells[1].text = resultado
    row_cells[2].text = tempo
    row_cells[3].text = cpu
    row_cells[4].text = memoria

# Salvar o arquivo
output_path = './doc.docx'
doc.save(output_path)

output_path
